from __future__ import annotations

import re
from pathlib import Path
from typing import List, Optional

from .models import MeetingMinutes, AgendaItem, TodoItem


def export_markdown(minutes: MeetingMinutes) -> str:
    lines = []

    lines.append(f"# {minutes.title or '会议纪要'}")
    lines.append("")

    if minutes.date:
        lines.append(f"**日期**：{minutes.date}")
    if minutes.attendees:
        lines.append(f"**参会人**：{', '.join(minutes.attendees)}")
    speakers = minutes.speakers()
    if speakers:
        lines.append(f"**发言人**：{', '.join(speakers)}")
    if minutes.source_file:
        lines.append(f"**来源文件**：{minutes.source_file}")
    lines.append("")

    lines.append("---")
    lines.append("")

    if minutes.agendas:
        lines.append("## 议题")
        lines.append("")
        for i, agenda in enumerate(minutes.agendas, 1):
            title = agenda.title
            source_note = ""
            if agenda.meeting_source:
                source_note = f" *[来源：{agenda.meeting_source}]*"
            lines.append(f"### {i}. {title}{source_note}")
            lines.append("")
            if agenda.content:
                lines.append(agenda.content)
                lines.append("")
            if agenda.key_sentences:
                lines.append("**要点**：")
                for ks in agenda.key_sentences:
                    lines.append(f"- {ks}")
                lines.append("")
            if agenda.risks:
                lines.append("**风险**：")
                for r in agenda.risks:
                    lines.append(f"- ⚠️ {r}")
                lines.append("")

    if minutes.highlights:
        lines.append("## 重点")
        lines.append("")
        for h in minutes.highlights:
            lines.append(f"- 🔑 {h}")
        lines.append("")

    if minutes.risks:
        lines.append("## 风险清单")
        lines.append("")
        for r in minutes.risks:
            lines.append(f"- ⚠️ {r}")
        lines.append("")

    if minutes.todos:
        lines.append("## 待办事项")
        lines.append("")
        has_meeting_source = any(t.get_all_sources() for t in minutes.todos)
        has_multiple = any(len(t.get_all_assignees()) > 1 or len(t.get_all_deadlines()) > 1 for t in minutes.todos)

        header_parts = ["#", "任务", "责任人", "截止时间", "优先级"]
        if has_multiple:
            header_parts = ["#", "任务", "责任人", "截止时间", "优先级"]
        if has_meeting_source:
            header_parts.append("会议来源")

        lines.append("| " + " | ".join(header_parts) + " |")
        lines.append("|" + "|".join(["---"] * len(header_parts)) + "|")
        for i, todo in enumerate(minutes.todos, 1):
            assignee = todo.display_assignees() or "❌ 未指定"
            deadline = todo.display_deadlines() or "❌ 未指定"
            priority = todo.priority or "-"
            row = [str(i), todo.task[:60].replace("|", "｜"), assignee, deadline, priority]
            if has_meeting_source:
                srcs = todo.get_all_sources()
                row.append("、".join(srcs) if srcs else "-")
            lines.append("| " + " | ".join(row) + " |")

        extra_notes = []
        for i, todo in enumerate(minutes.todos, 1):
            parts = []
            ag_src = todo.agenda_source or ""
            seg_src = todo.segment_source or ""
            if ag_src:
                parts.append(f"议题：{ag_src}")
            if seg_src:
                parts.append(f"发言：{seg_src}")
            if parts:
                extra_notes.append(f"- 待办#{i} 关联：{' | '.join(parts)}")
        if extra_notes:
            lines.append("")
            lines.append("**关联信息**：")
            for note in extra_notes:
                lines.append(note)
        lines.append("")

    if minutes.segments:
        lines.append("## 发言记录")
        lines.append("")
        for seg in minutes.segments:
            ts = f" [{seg.timestamp}]" if seg.timestamp else ""
            src = f" *[{seg.meeting_source}]*" if seg.meeting_source else ""
            lines.append(f"**{seg.speaker}**{ts}{src}：{seg.content}")
            lines.append("")

    lines.append("---")
    lines.append(f"*由 meeting-minutes 工具自动生成*")

    return "\n".join(lines)


def export_word(minutes: MeetingMinutes, output_path: Path) -> Path:
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        raise ImportError("导出 Word 需要安装 python-docx：pip install python-docx")

    doc = Document()

    title_para = doc.add_heading(minutes.title or "会议纪要", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if minutes.date:
        doc.add_paragraph(f"日期：{minutes.date}")
    if minutes.attendees:
        doc.add_paragraph(f"参会人：{', '.join(minutes.attendees)}")
    speakers = minutes.speakers()
    if speakers:
        doc.add_paragraph(f"发言人：{', '.join(speakers)}")

    doc.add_paragraph("─" * 40)

    if minutes.agendas:
        doc.add_heading("议题", level=1)
        for i, agenda in enumerate(minutes.agendas, 1):
            title_text = agenda.title
            if agenda.meeting_source:
                title_text += f"  [来源：{agenda.meeting_source}]"
            doc.add_heading(f"{i}. {title_text}", level=2)
            if agenda.content:
                doc.add_paragraph(agenda.content)
            if agenda.key_sentences:
                p = doc.add_paragraph()
                run = p.add_run("要点：")
                run.bold = True
                for ks in agenda.key_sentences:
                    doc.add_paragraph(ks, style="List Bullet")
            if agenda.risks:
                p = doc.add_paragraph()
                run = p.add_run("风险：")
                run.bold = True
                for r in agenda.risks:
                    doc.add_paragraph(f"⚠️ {r}", style="List Bullet")

    if minutes.highlights:
        doc.add_heading("重点", level=1)
        for h in minutes.highlights:
            doc.add_paragraph(f"🔑 {h}", style="List Bullet")

    if minutes.risks:
        doc.add_heading("风险清单", level=1)
        for r in minutes.risks:
            doc.add_paragraph(f"⚠️ {r}", style="List Bullet")

    if minutes.todos:
        doc.add_heading("待办事项", level=1)
        has_meeting_source = any(t.get_all_sources() for t in minutes.todos)
        cols = 6 if has_meeting_source else 5
        table = doc.add_table(rows=1, cols=cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["#", "任务", "责任人", "截止时间", "优先级"]
        if has_meeting_source:
            headers.append("会议来源")
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for idx, todo in enumerate(minutes.todos, 1):
            row = table.add_row()
            row.cells[0].text = str(idx)
            row.cells[1].text = todo.task[:60]
            row.cells[2].text = todo.display_assignees() or "❌ 未指定"
            row.cells[3].text = todo.display_deadlines() or "❌ 未指定"
            row.cells[4].text = todo.priority or "-"
            if has_meeting_source:
                srcs = todo.get_all_sources()
                row.cells[5].text = "、".join(srcs) if srcs else "-"
            missing = not todo.display_assignees() or not todo.display_deadlines()
            if missing:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 0, 0)

        extra_notes = []
        for i, todo in enumerate(minutes.todos, 1):
            parts = []
            if todo.agenda_source:
                parts.append(f"议题：{todo.agenda_source}")
            if todo.segment_source:
                parts.append(f"发言：{todo.segment_source}")
            if parts:
                extra_notes.append(f"待办#{i} 关联：{' | '.join(parts)}")
        if extra_notes:
            p = doc.add_paragraph()
            run = p.add_run("关联信息：")
            run.bold = True
            for note in extra_notes:
                doc.add_paragraph(note, style="List Bullet")

    if minutes.segments:
        doc.add_heading("发言记录", level=1)
        for seg in minutes.segments:
            p = doc.add_paragraph()
            run = p.add_run(f"{seg.speaker}")
            run.bold = True
            ts = f" [{seg.timestamp}]" if seg.timestamp else ""
            src = f" [{seg.meeting_source}]" if seg.meeting_source else ""
            p.add_run(f"{ts}{src}：{seg.content}")

    doc.add_paragraph("─" * 40)
    p = doc.add_paragraph("由 meeting-minutes 工具自动生成")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def export_minutes(minutes: MeetingMinutes, output_path: Path, fmt: str = "markdown") -> Path:
    fmt = fmt.lower().strip(".")
    if fmt in ("md", "markdown"):
        content = export_markdown(minutes)
        output_path = output_path.with_suffix(".md")
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(content, encoding="utf-8")
        return output_path
    elif fmt in ("docx", "word"):
        output_path = output_path.with_suffix(".docx")
        return export_word(minutes, output_path)
    else:
        raise ValueError(f"不支持的导出格式：{fmt}（支持 markdown/docx）")
